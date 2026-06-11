"""
continuidad_5A.py
Genera el documento de Continuidad Pedagógica — Historia 5° A (EES N°63, 2026)
Dependencia: pip install python-docx
Uso:         python continuidad_5A.py
Salida:      continuidad_pedagogica_5A.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ──────────────────────────────────────────────────────────────────
BLUE   = RGBColor(0x1F, 0x4E, 0x79)
LB     = RGBColor(0xD6, 0xE4, 0xF0)   # azul claro
MB     = RGBColor(0xAE, 0xD6, 0xF1)   # azul medio
YEL    = RGBColor(0xFF, 0xF9, 0xC4)   # amarillo
GRN    = RGBColor(0xE8, 0xF5, 0xE9)   # verde
PUR    = RGBColor(0xF3, 0xE5, 0xF5)   # violeta
WHT    = RGBColor(0xFF, 0xFF, 0xFF)
GRAY   = RGBColor(0x55, 0x55, 0x55)


# ── Helpers XML ──────────────────────────────────────────────────────────────
def hex_color(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color(rgb))
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def set_row_height(row, height_twips: int):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(height_twips))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


def set_table_borders(table, color="AAAAAA"):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_col_width(table, col_widths_twips):
    tbl = table._tbl
    tblGrid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    # insert after tblPr
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblPr.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)


def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.insert(0, tcW)


# ── Funciones de contenido ────────────────────────────────────────────────────
PAGE_WIDTH = 8496   # A4 con márgenes de 1.5 cm c/lado aprox (en twips: ~9026 con 1" margins)
                    # usamos 8496 = A4 content width con márgenes de 1.5cm

def add_paragraph(doc, text="", bold=False, italic=False, size=10,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=60, space_after=60, font="Arial"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Twips(space_before)
    p.paragraph_format.space_after = Twips(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = font
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def add_spacer(doc, before=80, after=80):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(before)
    p.paragraph_format.space_after = Twips(after)
    return p


def add_name_row(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(40)
    p.paragraph_format.space_after = Twips(40)
    run1 = p.add_run("Nombre y apellido: ")
    run1.bold = True
    run1.font.name = "Arial"
    run1.font.size = Pt(10)
    run2 = p.add_run("___________________________________________   Fecha: _______________")
    run2.font.name = "Arial"
    run2.font.size = Pt(10)


def full_width_table(doc, paragraphs_list, bg_color: RGBColor):
    """Tabla de una columna full-width con fondo de color."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, "AAAAAA")
    set_col_width(table, [PAGE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    set_cell_margins(cell)
    set_cell_width(cell, PAGE_WIDTH)
    # limpiar párrafo vacío por defecto
    cell._tc.clear_content()
    for para in paragraphs_list:
        cell._tc.append(para._p)
    return table


def header_box(doc, text):
    p = doc.add_paragraph()  # párrafo auxiliar
    p.paragraph_format.space_before = Twips(60)
    p.paragraph_format.space_after = Twips(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = WHT
    full_width_table(doc, [p], BLUE)
    # eliminar el párrafo suelto
    p._element.getparent().remove(p._element)


def section_header(doc, text, bg=None):
    if bg is None:
        bg = LB
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(60)
    p.paragraph_format.space_after = Twips(60)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    full_width_table(doc, [p], bg)
    p._element.getparent().remove(p._element)


def note_box(doc, text, bg=None):
    if bg is None:
        bg = YEL
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(60)
    p.paragraph_format.space_after = Twips(60)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    full_width_table(doc, [p], bg)
    p._element.getparent().remove(p._element)


def act_box(doc, num, title, ref):
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Twips(60)
    p1.paragraph_format.space_after = Twips(20)
    run1 = p1.add_run(f"ACTIVIDAD {num}  —  {title}")
    run1.bold = True
    run1.font.name = "Arial"
    run1.font.size = Pt(11)
    run1.font.color.rgb = WHT

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Twips(0)
    p2.paragraph_format.space_after = Twips(60)
    run2 = p2.add_run(f"Referencia en el módulo: {ref}")
    run2.italic = True
    run2.font.name = "Arial"
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    full_width_table(doc, [p1, p2], BLUE)
    for p in [p1, p2]:
        if p._element.getparent() is not None:
            p._element.getparent().remove(p._element)


def question(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Twips(100)
    p.paragraph_format.space_after = Twips(40)
    run_num = p.add_run(f"{num}. ")
    run_num.bold = True
    run_num.font.name = "Arial"
    run_num.font.size = Pt(10)
    run_q = p.add_run(text)
    run_q.font.name = "Arial"
    run_q.font.size = Pt(10)
    return p


def answer_box(doc, height_cm=3.5):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table, "AAAAAA")
    set_col_width(table, [PAGE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_margins(cell)
    set_cell_width(cell, PAGE_WIDTH)
    set_row_height(table.rows[0], int(height_cm * 567))  # 1cm ≈ 567 twips
    return table


def two_col_table(doc, w1, w2, rows):
    """rows = list of (col1_text, col2_text, is_header)"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_borders(table, "AAAAAA")
    set_col_width(table, [w1, w2])
    for i, (c1, c2, is_h) in enumerate(rows):
        row = table.rows[i]
        for j, (txt, w) in enumerate([(c1, w1), (c2, w2)]):
            cell = row.cells[j]
            set_cell_bg(cell, LB if is_h else WHT)
            set_cell_margins(cell)
            set_cell_width(cell, w)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = is_h
            run.font.name = "Arial"
            run.font.size = Pt(10)
            run.font.color.rgb = BLUE if is_h else RGBColor(0, 0, 0)
    return table


def bullet_item(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Twips(40)
    p.paragraph_format.space_after = Twips(40)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # Márgenes de página (A4, 1.5 cm)
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── ENCABEZADO ──────────────────────────────────────────────────────────
    header_box(doc, "DGCyE — Escuela de Educación Secundaria N° 63")
    add_spacer(doc, 10, 10)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Twips(60)
    p_title.paragraph_format.space_after = Twips(60)
    run = p_title.add_run("Historia — 5° A — Ciclo Lectivo 2026")
    run.bold = True; run.font.name = "Arial"; run.font.size = Pt(18)

    section_header(doc, "ACTIVIDADES DE CONTINUIDAD PEDAGÓGICA")
    add_paragraph(doc,
        "Estas actividades complementan y profundizan los contenidos del módulo. "
        "No repiten las guías de estudio: buscan que puedas relacionar, reflexionar y aplicar lo aprendido.",
        italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=40, space_after=40)

    add_spacer(doc, 40, 30)
    add_name_row(doc)
    add_spacer(doc, 20, 20)
    note_box(doc,
        "Leé con atención cada consigna antes de responder. "
        "Podés consultar tu carpeta y el módulo. "
        "Entregá las actividades a la docente al regreso.")
    add_spacer(doc, 60, 40)

    # ── ÍNDICE ───────────────────────────────────────────────────────────────
    section_header(doc, "ÍNDICE DE ACTIVIDADES")
    add_spacer(doc, 10, 10)
    actividades = [
        "Actividad 1  —  La Guerra Fría: el mundo bipolar desde adentro",
        "Actividad 2  —  La Guerra Fría en América Latina: intervención y soberanía",
        "Actividad 3  —  El peronismo: justicia social, derechos y conflicto",
        "Actividad 4  —  Antiperonismo y resistencia: proscripción e identidad",
        "Actividad 5  —  Argentina 1966–1976: autoritarismo, movilización y violencia",
        "Actividad 6  —  ESI: identidad, derechos y vida cotidiana (reflexión personal)",
    ]
    for a in actividades:
        bullet_item(doc, a)
    add_spacer(doc, 80, 80)

    # ════════════════════════════════════════════════════════════════════════
    # ACTIVIDAD 1
    # ════════════════════════════════════════════════════════════════════════
    act_box(doc, "1", "La Guerra Fría: el mundo bipolar desde adentro",
            "Capítulo 01 — págs. 10 a 19")
    add_spacer(doc, 20, 20)
    note_box(doc,
        "En el módulo estudiaste cómo funcionaban los dos bloques. "
        "Estas consignas te proponen ponerte en el lugar de quienes vivían ese mundo "
        "y pensar más allá de los datos.")
    add_spacer(doc, 30, 20)

    section_header(doc, "A. Ponerse en los zapatos del otro")
    add_spacer(doc, 10, 10)
    question(doc, "1",
        "Imaginá que sos un joven de 18 años en 1955. Elegí si vivís en Estados Unidos o en la "
        "Unión Soviética y explicá: ¿qué te decía tu gobierno sobre el otro bloque? ¿Qué información "
        "tenías y cuál te ocultaban? Usá ejemplos concretos del módulo.")
    answer_box(doc, 4.5)
    add_spacer(doc, 20, 10)
    question(doc, "2",
        "El módulo menciona que la Guerra Fría tuvo una fuerte \"connotación ideológica\". "
        "¿Qué significa eso? ¿Por qué no alcanzaba con tener más armas o más territorio "
        "para ganar ese conflicto?")
    answer_box(doc, 3.5)
    add_spacer(doc, 30, 20)

    section_header(doc, "B. Relaciones y contradicciones")
    add_spacer(doc, 10, 10)
    question(doc, "3",
        "Los Estados Unidos apoyaron dictaduras en nombre de la democracia. "
        "La URSS habló de libertad pero reprimió las rebeliones en sus países satélite. "
        "¿Cómo explicarías esa contradicción? ¿Qué nos dice sobre la diferencia entre "
        "el discurso y la práctica política?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "4",
        "La carrera armamentista y la carrera espacial fueron dos formas de competir sin "
        "disparar un tiro. ¿Por qué era tan importante ganar en esos terrenos? ¿Qué buscaba "
        "demostrar cada potencia ante su propio pueblo y ante el mundo?")
    answer_box(doc, 3.5)
    add_spacer(doc, 30, 20)

    section_header(doc, "C. Conexión con el presente", bg=GRN)
    add_spacer(doc, 10, 10)
    note_box(doc, "Para pensar libremente. No hay una única respuesta válida.", bg=GRN)
    add_spacer(doc, 10, 10)
    question(doc, "5",
        "¿Creés que hoy existe algo parecido a la lógica de bloques de la Guerra Fría? "
        "¿Qué países o conflictos actuales te recuerdan a ese período? "
        "Justificá con al menos un argumento concreto.")
    answer_box(doc, 4.0)
    add_spacer(doc, 100, 80)

    # ════════════════════════════════════════════════════════════════════════
    # ACTIVIDAD 2
    # ════════════════════════════════════════════════════════════════════════
    act_box(doc, "2", "La Guerra Fría en América Latina: intervención y soberanía",
            "Capítulo 04 — págs. 60 a 71")
    add_spacer(doc, 20, 20)
    note_box(doc,
        "América Latina fue uno de los escenarios más calientes de la Guerra Fría. "
        "Estas consignas te invitan a pensar cómo ese conflicto global afectó la vida "
        "concreta de los pueblos latinoamericanos.")
    add_spacer(doc, 30, 20)

    section_header(doc, "A. La lógica de la intervención")
    add_spacer(doc, 10, 10)
    question(doc, "1",
        "La Doctrina de Seguridad Nacional planteaba que los militares debían vigilar las "
        "\"fronteras ideológicas\" al interior de cada país, no solo las fronteras territoriales. "
        "¿Qué implicaba eso para los derechos de los ciudadanos? ¿Quiénes quedaban "
        "definidos como \"enemigos\"?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "2",
        "Los Estados Unidos apoyaron golpes de Estado en varios países latinoamericanos "
        "argumentando que defendían la democracia. ¿Qué contradicción ves en ese argumento? "
        "Explicá tu postura con ejemplos del módulo.")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "B. La Revolución cubana como punto de quiebre")
    add_spacer(doc, 10, 10)
    question(doc, "3",
        "La Revolución cubana fue un símbolo para muchos jóvenes latinoamericanos y una amenaza "
        "para los Estados Unidos. ¿Por qué el mismo hecho podía interpretarse de maneras tan "
        "opuestas? ¿Qué dependía del punto de vista de cada uno?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "4",
        "El módulo explica que Cuba pasó de ser un \"país semicolonial\" a aliarse con la URSS. "
        "¿Qué opciones tenía Cuba en el contexto de la Guerra Fría? ¿Podría haber tomado otro "
        "camino? Argumentá tu respuesta.")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "C. Para comparar")
    add_spacer(doc, 10, 10)
    question(doc, "5", "Completá el cuadro comparando la posición de ambas potencias en América Latina:")
    add_spacer(doc, 10, 10)
    two_col_table(doc, PAGE_WIDTH // 2, PAGE_WIDTH // 2, [
        ("Estados Unidos", "Unión Soviética", True),
        ("Objetivo principal en la región:", "Objetivo principal en la región:", False),
        ("Métodos que utilizó:", "Métodos que utilizó:", False),
        ("Ejemplo concreto del módulo:", "Ejemplo concreto del módulo:", False),
    ])
    add_spacer(doc, 100, 80)

    # ════════════════════════════════════════════════════════════════════════
    # ACTIVIDAD 3
    # ════════════════════════════════════════════════════════════════════════
    act_box(doc, "3", "El peronismo: justicia social, derechos y conflicto",
            "Capítulo 05 — págs. 74 a 81")
    add_spacer(doc, 20, 20)
    note_box(doc,
        "El peronismo transformó profundamente la sociedad argentina. "
        "Estas consignas buscan que vayas más allá de los hechos y analices "
        "sus sentidos, sus tensiones y sus consecuencias.")
    add_spacer(doc, 30, 20)

    section_header(doc, "A. Justicia social: ¿qué cambió y para quién?")
    add_spacer(doc, 10, 10)
    question(doc, "1",
        "El módulo describe mejoras concretas en las condiciones de vida de los trabajadores: "
        "salarios, vacaciones, vivienda, salud. ¿Por qué esos cambios generaban tanto entusiasmo "
        "en unos y tanta resistencia en otros? ¿Qué intereses concretos estaban en juego?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "2",
        "Perón decía que su modelo no era ni capitalismo puro ni comunismo, sino una \"tercera "
        "posición\". ¿Qué rol le asignaba al Estado en la economía? ¿En qué se diferenciaba eso "
        "del modelo soviético y del modelo liberal?")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "B. Tensiones y límites")
    add_spacer(doc, 10, 10)
    question(doc, "3",
        "El peronismo amplió derechos pero también concentró poder, limitó la prensa y persiguió "
        "opositores. ¿Cómo convivían esas dos caras en el mismo gobierno? ¿Es posible que un "
        "gobierno sea al mismo tiempo popular y autoritario? Argumentá con ejemplos.")
    answer_box(doc, 4.5)
    add_spacer(doc, 20, 10)
    question(doc, "4",
        "El voto femenino se conquistó en 1947. ¿Por qué ese derecho no existía antes? "
        "¿Qué cambió en la sociedad argentina a partir de esa ley? ¿Fue solo un acto de justicia "
        "o también una estrategia política? Podés sostener más de una posición.")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "C. Para reflexionar", bg=GRN)
    add_spacer(doc, 10, 10)
    note_box(doc, "Reflexión abierta. No hay una única respuesta válida.", bg=GRN)
    add_spacer(doc, 10, 10)
    question(doc, "5",
        "\"Sin desarrollo nacional no hay bienestar ni progreso.\" ¿Estás de acuerdo con esa "
        "afirmación? ¿Creés que hoy sigue siendo válida en la Argentina? "
        "Justificá con argumentos propios.")
    answer_box(doc, 4.0)
    add_spacer(doc, 100, 80)

    # ════════════════════════════════════════════════════════════════════════
    # ACTIVIDAD 4
    # ════════════════════════════════════════════════════════════════════════
    act_box(doc, "4", "Antiperonismo y resistencia: proscripción e identidad",
            "Capítulo 06 — págs. 84 a 97")
    add_spacer(doc, 20, 20)
    note_box(doc,
        "La caída de Perón no fue el fin del peronismo: fue el comienzo de una etapa marcada "
        "por la proscripción, la resistencia y una profunda división de la sociedad argentina.")
    add_spacer(doc, 30, 20)

    section_header(doc, "A. Prohibir una identidad")
    add_spacer(doc, 10, 10)
    question(doc, "1",
        "El Decreto 4161 prohibía mencionar el nombre de Perón, cantar la marcha peronista "
        "o exhibir sus imágenes. ¿Qué efecto buscaba producir esa prohibición en la sociedad? "
        "¿Por qué creés que no logró su objetivo? ¿Qué pasa cuando se intenta borrar "
        "una identidad política?")
    answer_box(doc, 4.5)
    add_spacer(doc, 20, 10)
    question(doc, "2",
        "El módulo describe formas de resistencia obrera: sabotajes, trabajo a desgano, comandos "
        "civiles. ¿Por qué surgieron esas formas de acción \"desde abajo\"? ¿Qué nos dicen sobre "
        "cómo la gente ejerce poder cuando los canales legales están cerrados?")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "B. Democracia condicionada")
    add_spacer(doc, 10, 10)
    question(doc, "3",
        "Entre 1955 y 1966, Argentina tuvo elecciones pero el peronismo estaba proscripto. "
        "¿Puede llamarse democracia a un sistema que excluye a la fuerza política mayoritaria? "
        "¿Qué condiciones considerás necesarias para que una democracia sea genuina?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "4",
        "Arturo Illia ganó con el 25% de los votos en un contexto de proscripción y fue "
        "derrocado antes de terminar su mandato. ¿Qué nos dice ese ciclo sobre la relación "
        "entre civiles y militares en esa época? ¿Quién tenía realmente el poder?")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "C. Línea de tiempo comentada")
    add_spacer(doc, 10, 10)
    note_box(doc,
        "Completá la siguiente tabla con los hechos clave del período 1955–1966. "
        "Para cada uno, agregá una breve explicación de su importancia histórica.")
    add_spacer(doc, 10, 10)
    w1 = int(PAGE_WIDTH * 0.30)
    w2 = PAGE_WIDTH - w1
    two_col_table(doc, w1, w2, [
        ("Año / Hecho", "¿Por qué fue importante?", True),
        ("1955 — Golpe contra Perón", "", False),
        ("1956 — Decreto 4161", "", False),
        ("1956 — Fusilamientos de José León Suárez", "", False),
        ("1958 — Pacto Perón-Frondizi", "", False),
        ("1963 — Gobierno de Illia", "", False),
    ])
    add_spacer(doc, 100, 80)

    # ════════════════════════════════════════════════════════════════════════
    # ACTIVIDAD 5
    # ════════════════════════════════════════════════════════════════════════
    act_box(doc, "5", "Argentina 1966–1976: autoritarismo, movilización y violencia",
            "Capítulo 09 — págs. 128 a 139")
    add_spacer(doc, 20, 20)
    note_box(doc,
        "Este período concentra algunas de las tensiones más fuertes de la historia argentina: "
        "dictadura, movilización popular, organizaciones armadas y el regreso de Perón. "
        "Estas consignas te proponen analizar esas tensiones en profundidad.")
    add_spacer(doc, 30, 20)

    section_header(doc, "A. El Cordobazo como hito")
    add_spacer(doc, 10, 10)
    question(doc, "1",
        "El Cordobazo fue una rebelión en la que obreros y estudiantes actuaron juntos. "
        "¿Por qué fue tan significativa esa alianza? ¿Qué tenían en común sus reclamos? "
        "¿En qué se diferencia una rebelión popular de una huelga común?")
    answer_box(doc, 4.5)
    add_spacer(doc, 20, 10)
    question(doc, "2",
        "El gobierno de Onganía prohibió los partidos políticos, intervino las universidades "
        "y censuró la cultura. ¿Qué modelo de sociedad buscaba construir? ¿Por qué esas medidas "
        "generaron más resistencia en lugar de más orden?")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "B. El regreso de Perón y la violencia política")
    add_spacer(doc, 10, 10)
    question(doc, "3",
        "El lema \"Cámpora al gobierno, Perón al poder\" resumía una situación muy particular. "
        "¿Qué significaba? ¿Por qué Perón no podía presentarse directamente a elecciones? "
        "¿Qué nos dice eso sobre los límites de la democracia en ese momento?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "4",
        "Las organizaciones armadas surgieron en un contexto de dictadura y proscripción. "
        "¿Qué argumentos usaban para justificar la lucha armada? ¿Estás de acuerdo con esa "
        "justificación? ¿Cuándo, si alguna vez, puede justificarse la violencia política? "
        "Argumentá con cuidado.")
    answer_box(doc, 4.5)
    add_spacer(doc, 30, 20)

    section_header(doc, "C. Para relacionar")
    add_spacer(doc, 10, 10)
    question(doc, "5",
        "En los años '60 y '70, Argentina vivió una espiral de violencia política. ¿Cómo influyó "
        "el contexto de la Guerra Fría en esa violencia? ¿En qué medida lo que pasaba en Argentina "
        "era parte de un fenómeno más amplio en América Latina?")
    answer_box(doc, 4.5)
    add_spacer(doc, 100, 80)

    # ════════════════════════════════════════════════════════════════════════
    # ACTIVIDAD 6
    # ════════════════════════════════════════════════════════════════════════
    act_box(doc, "6", "ESI: identidad, derechos y vida cotidiana",
            "Reflexión personal — no requiere consultar el módulo")
    add_spacer(doc, 20, 20)
    note_box(doc,
        "Esta actividad no evalúa conocimiento del módulo. Es un espacio para pensar, opinar "
        "y reflexionar desde tu propia experiencia y perspectiva. No hay respuestas correctas "
        "ni incorrectas: lo que importa es que argumentes con honestidad y profundidad.",
        bg=PUR)
    add_spacer(doc, 30, 20)

    section_header(doc, "A. Identidad y pertenencia")
    add_spacer(doc, 10, 10)
    question(doc, "1",
        "A lo largo de este año estudiamos cómo distintos grupos lucharon por ser reconocidos: "
        "obreros, mujeres, pueblos colonizados, peronistas proscriptos. ¿Con qué grupo o causa "
        "te identificaste más? ¿Por qué? ¿Encontrás alguna conexión con tu propia vida "
        "o con situaciones que conocés?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "2",
        "La identidad no es solo lo que somos: también es lo que nos dicen que somos, "
        "o lo que nos prohíben ser. ¿Alguna vez sentiste que algo de tu identidad —tu manera "
        "de ser, tu origen, tu género, tus gustos, tus creencias— fue cuestionado o "
        "invisibilizado? ¿Cómo lo viviste o cómo lo ves en otros?")
    answer_box(doc, 4.5)
    add_spacer(doc, 30, 20)

    section_header(doc, "B. Derechos: pasado y presente")
    add_spacer(doc, 10, 10)
    question(doc, "3",
        "Muchos derechos que hoy damos por sentados —el voto femenino, los derechos laborales, "
        "la libertad de expresión, los derechos de las personas LGBTQ+— fueron conquistados con "
        "lucha y muchas veces también arrebatados por la fuerza. ¿Qué derecho valorás más en tu "
        "vida cotidiana? ¿Qué pasaría si lo perdieras?")
    answer_box(doc, 4.0)
    add_spacer(doc, 20, 10)
    question(doc, "4",
        "Las Madres y Abuelas de Plaza de Mayo convirtieron el dolor personal en resistencia "
        "pública y colectiva. ¿Conocés algún ejemplo actual —en Argentina o en el mundo— de "
        "personas que transformen una injusticia personal en una causa colectiva? Contalo y "
        "explicá qué tiene en común con lo que hicieron ellas.")
    answer_box(doc, 4.0)
    add_spacer(doc, 30, 20)

    section_header(doc, "C. Memoria y futuro", bg=GRN)
    add_spacer(doc, 10, 10)
    note_box(doc, "Última reflexión del año. Tomá el tiempo que necesites.", bg=GRN)
    add_spacer(doc, 10, 10)
    question(doc, "5",
        "\"Para que no se repita\" es una de las frases más usadas cuando hablamos de la última "
        "dictadura. ¿Qué tendría que hacer concretamente una sociedad para que no se repita? "
        "¿Qué rol juegan la escuela, los medios de comunicación y vos mismo/a/e en eso?")
    answer_box(doc, 5.0)
    add_spacer(doc, 80, 60)

    # ── CRITERIOS DE EVALUACIÓN ──────────────────────────────────────────────
    section_header(doc, "CRITERIOS DE EVALUACIÓN")
    add_spacer(doc, 20, 20)
    w1 = int(PAGE_WIDTH * 0.37)
    w2 = PAGE_WIDTH - w1
    two_col_table(doc, w1, w2, [
        ("Criterio", "Se observará...", True),
        ("Comprensión de los temas",
         "Respuestas que demuestren manejo de los contenidos del módulo y la carpeta", False),
        ("Capacidad de relacionar",
         "Conexiones entre temas, períodos y escalas (Argentina / América Latina / mundo)", False),
        ("Reflexión crítica",
         "Opinión personal fundamentada con argumentos sólidos y coherentes", False),
        ("Expresión escrita",
         "Claridad, coherencia y uso de oraciones completas", False),
        ("Autonomía",
         "Uso pertinente del módulo y la carpeta como fuentes de consulta", False),
        ("Presentación",
         "Prolijidad, firma, fecha y entrega en tiempo y forma", False),
    ])
    add_spacer(doc, 60, 40)

    p_footer = add_paragraph(
        doc,
        "Docente: Vicente, Gabriela  —  Historia 5° A  —  EES N° 63  —  2026",
        italic=True, size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=60, space_after=40
    )

    return doc


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    output = "continuidad_pedagogica_5A.docx"
    doc = build_document()
    doc.save(output)
    print(f"Documento generado: {output}")
